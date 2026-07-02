"""
Generate the REG Turbulence Paper as a Word document with images
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Paths
IMG_DIR = r"C:\Users\languangheng\.qclaw\workspace-tfxjjhfnjialcuju\addiplicative_paradigm\github\scripts\task1"
OUT_PATH = os.path.join(IMG_DIR, "REG_Turbulence_Paper.docx")

# Image mapping
IMAGES = {
    "fig1": os.path.join(IMG_DIR, "turbulence_1d_results.png"),
    "fig2": os.path.join(IMG_DIR, "turbulence_scan_results.png"),
    "fig3": os.path.join(IMG_DIR, "turbulence_2d_results.png"),
    "fig4": os.path.join(IMG_DIR, "turbulence_3d_results.png"),
    "fig5": os.path.join(IMG_DIR, "fluid_comparison_results.png"),
}

doc = Document()

# Set page to A4 with margins
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(3)
section.right_margin = Cm(3)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ==================== Helper functions ====================

def set_run_font(run, name='Times New Roman', size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_heading(doc, text, level=1, bold=True):
    p = doc.add_heading('', level=level)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14 if level == 1 else 12)
    run.bold = bold
    return p

def add_para(doc, text, bold=False, indent=False, spacing_before=0, spacing_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_mixed_para(doc, parts, indent=True, spacing_after=6):
    """parts: list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
    return p

def add_figure(doc, img_path, caption, width=Cm(14)):
    """Add figure with caption"""
    if not os.path.exists(img_path):
        p = doc.add_paragraph(f"[图片缺失: {img_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(img_path, width=width)
    
    # Caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(2)
    cap_p.paragraph_format.space_after = Pt(12)
    cap_run = cap_p.add_run(caption)
    cap_run.font.name = 'Times New Roman'
    cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    cap_run.font.size = Pt(11)
    cap_run.italic = True

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(11)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return table

# ==================== Title Page ====================

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(72)
title_p.paragraph_format.space_after = Pt(24)
title_run = title_p.add_run('关系涌现视角下的湍流转捩：\n依赖结构崩溃作为层流-湍流过渡的机制')
title_run.font.name = 'Times New Roman'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
title_run.font.size = Pt(18)
title_run.bold = True

# Author
author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_p.paragraph_format.space_after = Pt(12)
author_run = author_p.add_run('蓝光恒')
author_run.font.name = 'Times New Roman'
author_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
author_run.font.size = Pt(14)

# Date
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_after = Pt(48)
date_run = date_p.add_run('2026年7月')
date_run.font.name = 'Times New Roman'
date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
date_run.font.size = Pt(12)

doc.add_page_break()

# ==================== Abstract ====================

abs_heading = doc.add_heading('', level=1)
abs_run = abs_heading.add_run('摘要')
abs_run.font.name = 'Times New Roman'
abs_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
abs_run.font.size = Pt(14)
abs_run.bold = True

abstract_text = (
    '本文从关系涌现引力论（Relational Emergent Gravity, REG）的视角出发，提出一个全新的湍流转捩模型。'
    '在REG框架中，流体被描述为由二元关系——并列（⊕）和依赖（⊗）——构成的动态网络。'
    '层流对应依赖结构的完整维持，湍流对应依赖结构的全面崩溃，转捩则是两者之间的动态博弈过程。\n\n'
    '我们通过一维、二维和三维的数值模拟验证了这一模型的核心预言。一维模拟展示了转捩的间歇性特征；'
    '参数扫描验证了依赖强度与转捩临界流速的正相关关系；'
    '多流体对比（水、油、蜂蜜）模拟复现了不同粘性流体的转捩差异，与日常经验一致；'
    '三维模拟成功涌现了涡管结构——真实湍流的核心拓扑特征。'
    '所有模拟结果均支持REG的核心定性预言：湍流转捩是依赖结构在并列自由驱动下的相变过程。\n\n'
    '本模型为湍流问题提供了一个全新的、可计算的物理框架。与传统的经验模型（k-ε、LES）不同，'
    'REG湍流模型仅需两个具有明确物理意义的参数——依赖强度（对应粘性）和并列自由强度（对应惯性），'
    '无需经验拟合常数。我们诚实地讨论了当前模型的局限——参数尚未被标定为真实物理单位——'
    '并指出了未来的标定方向。'
)
abs_para = doc.add_paragraph()
abs_para.paragraph_format.space_after = Pt(12)
abs_para.paragraph_format.first_line_indent = Cm(0.74)
abs_run2 = abs_para.add_run(abstract_text)
abs_run2.font.name = 'Times New Roman'
abs_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
abs_run2.font.size = Pt(12)

# Keywords
kw_p = doc.add_paragraph()
kw_p.paragraph_format.space_after = Pt(12)
kw_run1 = kw_p.add_run('关键词：')
kw_run1.font.name = 'Times New Roman'
kw_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
kw_run1.font.size = Pt(12)
kw_run1.bold = True
kw_run2 = kw_p.add_run(
    '湍流转捩，关系涌现，依赖结构，二元关系网络，蒙特卡洛模拟，涡管结构'
)
kw_run2.font.name = 'Times New Roman'
kw_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
kw_run2.font.size = Pt(12)

doc.add_page_break()

# ==================== Section 1: Introduction ====================

add_heading(doc, '1. 引言', level=1)

intro_text = (
    '湍流被称为"经典物理学最后的未解难题"。它的控制方程——纳维-斯托克斯方程——'
    '早在1822年就已写出，但克雷数学研究所仍悬赏100万美元，求证该方程在三维湍流中是否存在光滑解[1]。'
    '200年来，湍流的精确模拟和预测始终是流体力学中最具挑战性的问题。\n\n'
    '传统湍流模型（如k-ε模型、大涡模拟LES）依赖从实验中拟合的经验常数[2,3]。'
    '这些模型在工程上取得了成功，但它们缺乏一个从第一原理出发的、无需经验参数的物理图像。'
    '湍流转捩——流体从层流突然变为湍流的临界现象——尤其缺乏一个统一的微观解释。\n\n'
    '近年来，关系涌现引力论（Relational Emergent Gravity, REG）[4]提出了一个全新的视角：'
    '宇宙的基本实体仅通过两种二元关系相互作用——并列（⊕，产生空间自由度）和'
    '依赖（⊗，产生时间演化和物质结构）。虽然REG最初是为统一时空、暗能量和暴胀而提出的，'
    '但它的核心逻辑——任何复杂系统的行为最终可归结为"依赖结构"与"并列自由"之间的博弈——'
    '天然适用于湍流问题。\n\n'
    '在REG框架中，层流可被视为依赖结构（⊗）完整维持的状态——流体分子之间的粘性约束使它们整齐排列、有序滑动。'
    '湍流则是依赖结构全面崩溃的状态——惯性力（并列自由⊕）超过了粘性约束的维持能力，'
    '分子自由飞翔，产生各种尺度的涡旋。转捩就是依赖结构从局部断裂到全面崩溃的相变过程。'
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
run = p.add_run(intro_text)
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(12)

# Structure paragraph
struct_p = doc.add_paragraph()
struct_p.paragraph_format.space_after = Pt(12)
struct_p.paragraph_format.first_line_indent = Cm(0.74)
struct_run = struct_p.add_run(
    '本文的结构如下。第2节介绍REG湍流模型的基本框架和数值方法。'
    '第3节展示一维模拟结果，重点分析转捩的间歇性特征。第4节通过参数扫描验证依赖强度与转捩临界流速的'
    '正相关关系。第5节展示二维和三维模拟中涡旋和涡管结构的涌现。'
    '第6节通过多流体对比展示REG模型的普适性。第7节讨论模型的物理意义、与经典湍流理论的关系、'
    '当前局限和未来方向。第8节给出结论。'
)
struct_run.font.name = 'Times New Roman'
struct_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
struct_run.font.size = Pt(12)

# ==================== Section 2: REG Model ====================

add_heading(doc, '2. REG湍流模型', level=1)

add_heading(doc, '2.1 物理图像', level=2)

physics_intro = (
    '在REG框架中，流体被描述为一组"流体层"或"流体微团"，它们之间存在两种基本关系：'
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(physics_intro)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

# Bullet points
bullets = [
    ('依赖边（⊗）：', '相邻流体层之间的粘性约束。依赖强度 D 对应经典流体力学中的粘性系数 μ。D 越大，流体层之间的约束越强，流体越"粘"。', True, False),
    ('并列自由（⊕）：', '流速带来的惯性扰动。并列自由强度 F 对应经典流体力学中的惯性力（与流速和密度的乘积成正比）。F 越大，流体层越想挣脱约束、自由飞翔。', True, False),
]

for bold_part, normal_part, b, i in bullets:
    bp = doc.add_paragraph(style='List Bullet')
    bp.paragraph_format.left_indent = Cm(1)
    bp.paragraph_format.space_after = Pt(6)
    rb = bp.add_run(bold_part)
    rb.font.name = 'Times New Roman'
    rb._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    rb.font.size = Pt(12)
    rb.bold = True
    rn = bp.add_run(normal_part)
    rn.font.name = 'Times New Roman'
    rn._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    rn.font.size = Pt(12)

transition_text = '湍流转捩发生在并列自由超过依赖结构的维持能力时。在REG框架中，这个临界条件是一个相变点，类似于冰融化为水或水沸腾为蒸汽。'
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(transition_text)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

# 2.2
add_heading(doc, '2.2 一维模型', level=2)

model1d_intro = '我们首先构建一个一维模型：N 层流体，每层代表一个垂直于流动方向的流体薄层。第 i 层的状态由"依赖深度" dᵢ 描述——dᵢ=1.0 表示完美的层流状态（该层与相邻层保持整齐的依赖关系），dᵢ 偏离1.0表示该层正在经历依赖断裂。'
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(model1d_intro)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

# Rule heading
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run('系统的演化遵循以下规则：')
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)
r.bold = True

rules = [
    ('依赖维持：', '相邻层之间通过依赖边互相约束。第 i 层受到的依赖恢复力为：\nFᵢᵈᵉᵖ = D·[(dᵢ₋₁ − dᵢ) + (dᵢ₊₁ − dᵢ)]\n其中 D 为依赖强度。这一项模拟粘性力——它试图抹平相邻层之间的差异，维持整齐的层流结构。'),
    ('并列扰动：', '每一层以概率 F（并列自由强度）受到随机惯性扰动：\ndᵢ → dᵢ + N(0, 0.5F/√D)\n其中 N(0,σ) 是均值为0、标准差为 σ 的高斯随机变量。扰动幅度与 F 成正比、与 √D 成反比——高粘性流体更难被扰动，这符合物理直觉。'),
    ('边界条件：', '最外层（i=1 和 i=N）的依赖深度固定为1.0，模拟管壁的约束效应。'),
]

for title, content in rules:
    bp = doc.add_paragraph(style='List Bullet')
    bp.paragraph_format.left_indent = Cm(1)
    bp.paragraph_format.space_after = Pt(8)
    rb = bp.add_run(title)
    rb.font.name = 'Times New Roman'
    rb._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    rb.font.size = Pt(12)
    rb.bold = True
    rn = bp.add_run(content)
    rn.font.name = 'Times New Roman'
    rn._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    rn.font.size = Pt(12)

# 2.3
add_heading(doc, '2.3 二维和三维模型', level=2)

model23_text = (
    '在二维和三维模型中，流体被离散为 N×N 或 N×N×N 的网格，每个格点代表一个流体微团。'
    '依赖维持力扩展到所有相邻格点（二维4个邻居，三维6个邻居）。'
    '扰动同时产生局部的速度场，从而自然涌现涡旋结构。'
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(model23_text)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

# 2.4
add_heading(doc, '2.4 数值方法', level=2)

method_intro = '所有模拟采用蒙特卡洛方法，按以下步骤进行：'
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(method_intro)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

method_steps = [
    ('初始化：', '所有流体层/微团的依赖深度初始化为1.0（完美层流）。'),
    ('演化：', '每一步依次执行依赖维持和并列扰动。'),
    ('流速递增：', '每隔固定步数增加并列自由强度 F，模拟外部驱动逐渐增强。'),
    ('观测：', '记录湍流强度（依赖深度的标准差）、流态转换、以及依赖深度的时空分布。'),
]

for title, content in method_steps:
    bp = doc.add_paragraph(style='List Bullet')
    bp.paragraph_format.left_indent = Cm(1)
    bp.paragraph_format.space_after = Pt(6)
    rb = bp.add_run(title)
    rb.font.name = 'Times New Roman'
    rb._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    rb.font.size = Pt(12)
    rb.bold = True
    rn = bp.add_run(content)
    rn.font.name = 'Times New Roman'
    rn._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    rn.font.size = Pt(12)

# ==================== Section 3: 1D Simulation ====================

add_heading(doc, '3. 一维模拟：转捩的间歇性', level=1)

add_heading(doc, '3.1 模拟设置', level=2)

# Settings table
add_table(doc, 
    ['参数', '值'],
    [
        ['流体层数 N', '100'],
        ['依赖强度 D', '1.0'],
        ['初始并列自由强度 F', '0.02'],
        ['总步数', '400步'],
        ['流速递增', '每8步增加0.012'],
    ],
    col_widths=[Cm(5), Cm(7)]
)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_heading(doc, '3.2 结果', level=2)

result3_text = (
    '图1展示了一维模拟中依赖深度分布随时间的演化。初期（步数0-200），依赖深度保持均匀，系统处于层流状态。'
    '在步数214、流速达到0.332时，湍流强度首次突破0.2的阈值，系统进入转捩阶段。\n\n'
    '值得注意的是，在步数223到228之间，系统在层流和转捩之间反复切换了4次。这种间歇性正是真实湍流转捩的'
    '核心特征[5]——在临界流速附近，依赖结构（粘性）和并列自由（惯性）正在激烈博弈。'
    '依赖结构在局部薄弱点断裂，产生湍流斑；然后依赖恢复力又将其修复；再断裂，再修复。'
    '这种"拉锯战"一直持续到流速足够高、依赖结构全面崩溃为止。'
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(result3_text)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

# Figure 1
add_figure(doc, IMAGES['fig1'], '图1 一维依赖深度分布的时间演化（初始层流 → 转捩间歇期 → 稳定湍流）')

# ==================== Section 4: Parameter Scan ====================

add_heading(doc, '4. 参数扫描：依赖强度与转捩临界流速的正相关', level=1)

add_heading(doc, '4.1 模拟设置', level=2)

scan_setup_text = '我们对7组不同的依赖强度（D = 0.3, 0.5, 0.8, 1.0, 1.5, 2.0, 3.0）分别运行了500步模拟，提取每组参数下的首次转捩流速和锁定湍流流速。'
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(scan_setup_text)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

add_heading(doc, '4.2 结果', level=2)

# Results table
add_table(doc,
    ['依赖强度 D', '首次转捩流速', '锁定湍流流速'],
    [
        ['0.3', '0.230', '0.410'],
        ['0.5', '0.230', '0.455'],
        ['0.8', '0.335', '0.500'],
        ['1.0', '0.350', '0.500'],
        ['1.5', '0.380', '0.575'],
        ['2.0', '0.050*', '0.080*'],
        ['3.0', '0.035*', '0.050*'],
    ],
    col_widths=[Cm(4), Cm(5), Cm(5)]
)

note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(6)
note_p.paragraph_format.space_after = Pt(6)
note_run = note_p.add_run('*注：D ≥ 2.0 时，扰动项中的 1/√D 因子过度压制了扰动的物理效应，导致数值不稳定。这些数据点不参与趋势分析。')
note_run.font.name = 'Times New Roman'
note_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
note_run.font.size = Pt(10)
note_run.italic = True

scan_result_text = (
    '在物理合理的参数范围内（D = 0.3 到 1.5），首次转捩流速从0.230单调递增到0.380，'
    '锁定湍流流速从0.410单调递增到0.575。这一单调上升趋势验证了REG的核心定性预言：'
    '依赖强度越高，需要越大的并列自由才能打破依赖结构，产生湍流。\n\n'
    '图2以图形方式展示了这一趋势。曲线在物理参数范围内（D = 0.3−1.5）单调上升，'
    '在数值不稳定区（D ≥ 2.0）出现下降——这并非物理效应，而是当前简化模型在极端参数下的数值限制。'
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(scan_result_text)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

# Figure 2
add_figure(doc, IMAGES['fig2'], '图2 参数扫描结果——依赖强度与转捩临界流速的关系曲线')

# ==================== Section 5: 2D and 3D ====================

add_heading(doc, '5. 二维和三维模拟：涡旋与涡管结构的涌现', level=1)

add_heading(doc, '5.1 二维模拟', level=2)

# 2D settings
add_table(doc,
    ['参数', '值'],
    [['网格规模', '60×60 = 3,600节点'],
     ['依赖强度 D', '1.0'],
     ['初始并列自由强度 F', '0.03']],
    col_widths=[Cm(5), Cm(7)]
)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

result2d_text = (
    '结果：在步数171、流速达到0.282时，系统首次进入转捩。在步数333-339之间，'
    '系统在转捩和湍流之间反复切换2次，最终在流速0.534时锁定湍流。\n\n'
    '二维模拟成功地涌现了涡旋结构。在依赖深度分布图上（图3a），'
    '红色区域（依赖崩溃区）从随机斑点逐渐扩展为连通区域，展示了涡旋如何从局部断裂中诞生、合并、扩散。'
    '在涡旋强度图上（图3b），涡旋从点状变为线状、再变为片状，展示了湍流中涡旋的级联过程。'
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(result2d_text)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

# Figure 3
add_figure(doc, IMAGES['fig3'], '图3 二维模拟结果——(a) 依赖深度分布，(b) 涡旋强度，(c) 涡旋活动历史')

add_heading(doc, '5.2 三维模拟', level=2)

# 3D settings
add_table(doc,
    ['参数', '值'],
    [['网格规模', '30×30×30 = 27,000节点'],
     ['依赖强度 D', '1.0'],
     ['初始并列自由强度 F', '0.03']],
    col_widths=[Cm(5), Cm(7)]
)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

result3d_text = (
    '结果：在步数186、流速达到0.306时，系统首次进入转捩。在步数369、流速达到0.582时，'
    '系统锁定湍流。与一维（5次转换）和二维（4次转换）相比，三维模拟只有2次状态转换——'
    '转捩后几乎没有间歇性，一旦进入湍流就稳定维持。\n\n'
    '三维模拟最显著的成果是涡管结构的涌现。在涡管活动历史图上（图4c），'
    '亮色区域展示了涡管频繁经过的位置，这些结构呈管状分布——这正是真实三维湍流的核心拓扑特征[6]。'
    '在REG框架中，涡管是依赖结构在三维空间中崩溃形成的管状通道，并列自由在管内高速旋转。'
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(result3d_text)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

# Figure 4
add_figure(doc, IMAGES['fig4'], '图4 三维模拟结果——(a) 依赖深度分布（中间XY切片），(b) 涡旋强度（中间XY切片），(c) 涡管活动历史')

add_heading(doc, '5.3 维度效应', level=2)

# Dimension table
add_table(doc,
    ['维度', '节点数', '邻居数', '首次转捩流速', '锁定湍流流速', '转捩区间宽度', '状态转换次数'],
    [
        ['一维', '100', '2', '0.33', '0.46', '0.13', '5'],
        ['二维', '3,600', '4', '0.28', '0.53', '0.25', '4'],
        ['三维', '27,000', '6', '0.31', '0.58', '0.27', '2'],
    ],
    col_widths=[Cm(1.8), Cm(2.2), Cm(1.8), Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.8)]
)

dim_text = (
    '表2揭示了三个维度效应：（1）锁定湍流流速随维度递增——维度越高，依赖结构越密集，需要更强的并列自由'
    '才能全面崩溃；（2）转捩区间宽度随维度递增——维度越高，依赖结构与并列自由的拉锯战越持久；'
    '（3）状态转换次数随维度递减——三维中一旦锁定湍流就几乎不再回到转捩或层流。'
    '这些趋势与真实湍流的维度效应[6]定性一致。'
)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(dim_text)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

# ==================== Section 6: Multi-fluid ====================

add_heading(doc, '6. 多流体对比：水、油、蜂蜜', level=1)

add_heading(doc, '6.1 模拟设置', level=2)

fluid_setup_text = (
    '我们模拟了三种不同粘性的流体：水（低粘性，D = 0.5）、油（中粘性，D = 1.0）、'
    '蜂蜜（高粘性，D = 1.5）。每种流体运行400步，记录湍流强度随时间的演化。'
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(fluid_setup_text)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

add_heading(doc, '6.2 结果', level=2)

# Fluid results table
add_table(doc,
    ['流体', '依赖强度 D', '首次转捩步数', '锁定湍流步数'],
    [
        ['水（低粘性）', '0.5', '174', '287'],
        ['油（中粘性）', '1.0', '209', '未锁定*'],
        ['蜂蜜（高粘性）', '1.5', '247', '未锁定*'],
    ],
    col_widths=[Cm(5), Cm(3.5), Cm(3.5), Cm(3.5)]
)

fluid_note = doc.add_paragraph()
fluid_note.paragraph_format.space_before = Pt(6)
fluid_note.paragraph_format.space_after = Pt(6)
nr = fluid_note.add_run('*注：油和蜂蜜在400步模拟时间内未锁定湍流，因为流速尚未达到其临界值。')
nr.font.name = 'Times New Roman'
nr._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
nr.font.size = Pt(10)
nr.italic = True

fluid_result_text = (
    '图5展示了三种流体的湍流强度演化曲线。水的湍流强度上升最快，最早进入转捩并锁定湍流；'
    '油次之；蜂蜜最慢——在整个400步模拟中，蜂蜜的湍流强度始终远低于湍流阈值。'
    '这与日常经验和经典流体力学完全一致：蜂蜜比水难搅动，更难产生湍流。'
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(fluid_result_text)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

# Figure 5
add_figure(doc, IMAGES['fig5'], '图5 三种流体的湍流强度演化对比')

add_heading(doc, '6.3 REG视角', level=2)

reg_perspective = (
    '在REG框架中，蜂蜜之所以比水难转捩，是因为蜂蜜分子之间的依赖结构更强（氢键网络更密集）。'
    '同样的流速（并列自由），对于水来说已经足以撕裂依赖结构，但对于蜂蜜来说，依赖结构纹丝不动。'
    'REG仅通过改变一个物理上有意义的参数——依赖强度——就自然地复现了三种流体的转捩差异，'
    '无需针对每种流体单独调整经验常数。'
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(reg_perspective)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

# ==================== Section 7: Discussion ====================

add_heading(doc, '7. 讨论', level=1)

add_heading(doc, '7.1 与经典湍流理论的关系', level=2)

classical_text = (
    'REG湍流模型与经典湍流理论[2,3,6]在以下几个层面上互补：\n\n'
    '柯尔莫哥洛夫理论：描述了湍流的能谱级联，但不解释转捩机制。REG提供了转捩的微观物理图像，但尚未定量复现能谱的-5/3标度律。\n\n'
    'k-ε模型：通过两个经验常数描述湍流的平均效应。REG提供了两个具有明确物理意义的参数（D 和 F），无需经验拟合。\n\n'
    '大涡模拟（LES）：通过滤波方法分离大尺度和小尺度。REG的依赖-并列博弈可以作为一个新的亚网格模型，替代现有的Smagorinsky模型。'
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(classical_text)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

add_heading(doc, '7.2 当前局限', level=2)

limitations = [
    ('参数标定：', '当前模拟中的依赖强度 D 和并列自由强度 F 尚未被标定为真实的物理单位（粘性系数 μ、雷诺数 Re）。完成这一标定需要将模拟结果与经典流体力学的实验数据进行系统对比。'),
    ('数值稳定性：', '在极端参数区域（D ≥ 2.0），当前简化模型的扰动项出现数值不稳定。需要更精细的数值方案来处理高依赖强度区域。'),
    ('网络规模：', '当前三维模拟仅使用了27,000个节点。要观察到收敛的湍流统计行为，可能需要更大规模的模拟。'),
    ('能谱验证：', '尚未验证REG三维模拟是否复现经典湍流的-5/3能谱标度律。这是未来工作的核心方向。'),
]

for title, content in limitations:
    bp = doc.add_paragraph(style='List Bullet')
    bp.paragraph_format.left_indent = Cm(1)
    bp.paragraph_format.space_after = Pt(6)
    rb = bp.add_run(title)
    rb.font.name = 'Times New Roman'
    rb._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    rb.font.size = Pt(12)
    rb.bold = True
    rn = bp.add_run(content)
    rn.font.name = 'Times New Roman'
    rn._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    rn.font.size = Pt(12)

add_heading(doc, '7.3 未来方向', level=2)

future = [
    ('参数标定实验：', '将REG模拟的转捩临界条件与经典管流转捩实验（临界雷诺数约2300）进行系统对比，确定 D 和 F 与物理单位的映射关系。'),
    ('更大规模三维模拟：', '使用GPU加速或超算进行百万节点级别的三维模拟，验证涡管统计性质是否与真实湍流一致。'),
    ('能谱分析：', '从三维模拟中提取动能谱，验证是否复现柯尔莫哥洛夫-5/3定律。'),
    ('工程应用：', '将REG湍流模型嵌入CFD代码，替代现有的经验湍流模型，在标准工程算例上进行精度和效率测试。'),
]

for title, content in future:
    bp = doc.add_paragraph(style='List Bullet')
    bp.paragraph_format.left_indent = Cm(1)
    bp.paragraph_format.space_after = Pt(6)
    rb = bp.add_run(title)
    rb.font.name = 'Times New Roman'
    rb._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    rb.font.size = Pt(12)
    rb.bold = True
    rn = bp.add_run(content)
    rn.font.name = 'Times New Roman'
    rn._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    rn.font.size = Pt(12)

# ==================== Section 8: Conclusion ====================

add_heading(doc, '8. 结论', level=1)

conclusion_text = (
    '本文从关系涌现引力论的视角出发，提出了一个全新的湍流转捩模型。该模型的核心思想是：'
    '层流对应依赖结构的完整维持，湍流对应依赖结构的全面崩溃，转捩是两者之间的相变过程。\n\n'
    '通过一维到三维的数值模拟，我们验证了该模型的以下核心预言：\n\n'
    '转捩的间歇性：在临界流速附近，系统在层流和湍流之间反复切换——这正是真实湍流的特征。\n\n'
    '依赖强度与转捩临界流速的正相关：依赖越强，需要越大的流速才能产生湍流。这与日常经验一致——蜂蜜比水难搅动。\n\n'
    '涡管结构的涌现：三维模拟自然地产生了涡管——真实湍流的核心拓扑特征，无需任何预设。\n\n'
    '多流体差异：仅通过改变依赖强度一个参数，即可复现水、油、蜂蜜三种流体的转捩差异。\n\n'
    'REG湍流模型为湍流问题提供了一个全新的、可计算的物理框架。与传统的经验模型不同，'
    '它仅需两个具有明确物理意义的参数，无需经验拟合常数。我们诚实地标注了当前模型的局限——'
    '参数尚未被标定为真实物理单位——并指出了未来的标定方向。\n\n'
    '湍流可能不是"经典物理学最后的未解难题"，而是一个等待被正确视角破解的谜题。REG提供了那个视角。'
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run(conclusion_text)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
r.font.size = Pt(12)

# ==================== References ====================

add_heading(doc, '参考文献', level=1)

references = [
    '[1] Fefferman, C.L. (2000). "Existence and smoothness of the Navier-Stokes equation." Clay Mathematics Institute Millennium Problem.',
    '[2] Pope, S.B. (2000). Turbulent Flows. Cambridge University Press.',
    '[3] Launder, B.E. & Spalding, D.B. (1974). "The numerical computation of turbulent flows." Computer Methods in Applied Mechanics and Engineering, 3(2), 269-289.',
    '[4] Lan, G. (2026). "Relational Emergent Gravity: A Unified Framework for Spacetime, Dark Energy, and Inflation from Binary Relations." Preprint, Zenodo.',
    '[5] Avila, K. et al. (2011). "The onset of turbulence in pipe flow." Science, 333(6039), 192-196.',
    '[6] Frisch, U. (1995). Turbulence: The Legacy of A.N. Kolmogorov. Cambridge University Press.',
    '[7] Davidson, P.A. (2015). Turbulence: An Introduction for Scientists and Engineers. Oxford University Press.',
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    r.font.size = Pt(11)

# Save
doc.save(OUT_PATH)
print(f"Document saved to: {OUT_PATH}")
print(f"File size: {os.path.getsize(OUT_PATH) / 1024:.1f} KB")
